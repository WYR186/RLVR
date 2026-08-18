"""Render the zero-background Chinese results explainer to .docx.

Reuses the house style helpers from build_exp1_guide so both Chinese Word
documents look like one series. This document has no generated figures, so
the renderer here is the plain markdown subset the source file actually uses.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

import build_exp1_guide as house_style
from build_exp1_guide import (
    BLUE,
    DARK_BLUE,
    LIGHT_GRAY,
    MUTED,
    NAVY,
    add_code_block,
    add_field,
    add_inline_runs,
    add_numbering_definition,
    add_table,
    apply_num,
    configure_styles,
    parse_table,
    set_paragraph_border,
    set_paragraph_shading,
    set_run_font,
    set_update_fields,
)

# LibreOffice's headless macOS exporter drops CJK glyphs from the system TTC
# fonts. The QA command supplies a task-local Fontconfig entry for the
# TrueType Noto Sans SC shipped with an installed app; Word can also substitute
# this common CJK family cleanly on machines where it is not installed.
house_style.CJK_FONT = "Noto Sans SC"
house_style.WESTERN_FONT = "Arial"
house_style.MONO_FONT = "Menlo"
house_style.style_font.__kwdefaults__.update(
    western="Arial", east_asia="Noto Sans SC")
house_style.set_run_font.__kwdefaults__.update(
    western="Arial", east_asia="Noto Sans SC")


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "experiment 1" / "EXP1_结果详解_中文版.md"
OUTPUT = ROOT / "experiment 1" / "EXP1_结果详解_中文版.docx"

# Arial Unicode MS has no glyphs for these; Word renders them as boxes.
# The spaced variants come first so the label absorbs the space after the emoji.
EMOJI_REPLACEMENTS = {
    "⚠️ ": "注意：",
    "📌 ": "要点：",
    "⚠️": "注意：",
    "📌": "要点：",
    "①": "(1)",
    "②": "(2)",
}


def strip_unsupported_glyphs(text: str) -> str:
    for bad, good in EMOJI_REPLACEMENTS.items():
        text = text.replace(bad, good)
    return text


def configure_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(2)
    run = hp.add_run("实验一 结果详解 · 事实核验修订版")
    set_run_font(run, size=9, color=MUTED, bold=True)
    set_paragraph_border(hp, "bottom", "D7DBE2", size=5, space=3)

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(3)
    run = fp.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    add_field(fp, "PAGE", "1")
    run = fp.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)


def add_title_block(doc: Document, title: str) -> None:
    p = doc.add_paragraph(style="Title")
    add_inline_runs(p, title)
    # This revision has a longer, audit-specific title than the house-style
    # default; 24 pt keeps the full title balanced on the first page.
    for run in p.runs:
        run.font.size = Pt(20)

    p = doc.add_paragraph(style="Subtitle")
    add_inline_runs(p, "Algoverse · RLVR 可塑性项目 · Experiment 1 工件级事实核验")


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("目录")
    set_run_font(run, size=13, bold=True, color=NAVY)
    set_paragraph_border(p, "bottom", "D7DBE2", size=5, space=3)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    add_field(p, r'TOC \o "1-2" \h \z \u', "在 Word 中按 F9 更新目录")

    p = doc.add_paragraph()
    run = p.add_run("（在 Word 中打开后，右键目录 → 更新域，即可显示页码。）")
    set_run_font(run, size=9, color=MUTED, italic=True)


def join_wrapped(lines: list[str]) -> str:
    """Reflow source-wrapped lines into one string.

    Markdown line breaks inside a paragraph are wrapping, not content, so an
    inline span such as **...** may straddle them. Rejoining before the inline
    parser runs is what keeps those markers from leaking into the output.
    """
    text = ""
    for line in lines:
        if not text:
            text = line
        elif _is_latin_word_char(text[-1]) and _is_latin_word_char(line[0]):
            text += " " + line
        else:
            # Chinese needs no space at a wrap, and neither does a boundary
            # touching punctuation or an inline marker.
            text += line
    return text


def _is_latin_word_char(ch: str) -> bool:
    return ch.isascii() and ch.isalnum()


def add_callout(doc: Document, lines: list[str]) -> None:
    # Blank quote lines separate paragraphs inside one callout box.
    blocks: list[list[str]] = [[]]
    for line in lines:
        if line:
            blocks[-1].append(line)
        elif blocks[-1]:
            blocks.append([])
    blocks = [b for b in blocks if b]

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    set_paragraph_shading(p, LIGHT_GRAY)
    set_paragraph_border(p, "left", BLUE, size=14, space=5)
    for idx, block in enumerate(blocks):
        if idx:
            p.add_run().add_break()
        add_inline_runs(p, join_wrapped(block))


def build_body(doc: Document) -> None:
    raw = strip_unsupported_glyphs(SOURCE.read_text(encoding="utf-8"))
    lines = raw.splitlines()
    # The H1 title is rendered by add_title_block; the body starts after it.
    start = next(i for i, line in enumerate(lines) if line.startswith("# ")) + 1
    lines = lines[start:]

    i = 0
    bullet_num_id = add_numbering_definition(doc, ordered=False)
    current_ordered_num_id = None
    in_ordered_block = False

    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()
        indent = len(raw) - len(raw.lstrip())
        if not stripped:
            in_ordered_block = False
            i += 1
            continue

        if stripped == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            set_paragraph_border(p, "bottom", "D7DBE2", size=5, space=2)
            i += 1
            continue

        if stripped.startswith("```"):
            code = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i].rstrip())
                i += 1
            i += 1
            add_code_block(doc, code)
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            add_table(doc, parse_table(table_lines))
            in_ordered_block = False
            continue

        heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading:
            # The source's only H1 is the title, so shift H2 -> Heading 1.
            level = min(len(heading.group(1)) - 1, 3)
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline_runs(p, heading.group(2))
            i += 1
            continue

        if stripped.startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip().lstrip(">").strip())
                i += 1
            add_callout(doc, quote_lines)
            continue

        bullet = re.match(r"^-\s+(.+)$", stripped)
        ordered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if bullet or ordered:
            item_lines = [(bullet or ordered).group(1)]
            i += 1
            # A wrapped list item continues on lines indented past its marker.
            while i < len(lines):
                nxt = lines[i].strip()
                nxt_indent = len(lines[i]) - len(lines[i].lstrip())
                if not nxt or nxt_indent <= indent or re.match(r"^(-|\d+\.)\s+", nxt):
                    break
                item_lines.append(nxt)
                i += 1

            p = doc.add_paragraph()
            if bullet:
                apply_num(p, bullet_num_id)
            else:
                if not in_ordered_block:
                    current_ordered_num_id = add_numbering_definition(doc, ordered=True)
                    in_ordered_block = True
                apply_num(p, current_ordered_num_id)
            if indent:
                # The numbering definition indents level 0 by 0.375"; nested
                # items step in by one more level each.
                p.paragraph_format.left_indent = Inches(0.375 * (1 + indent // 2))
            add_inline_runs(p, join_wrapped(item_lines))
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (not nxt or nxt == "---" or nxt.startswith("#") or nxt.startswith("|") or
                    nxt.startswith("```") or nxt.startswith(">") or re.match(r"^-\s+", nxt) or
                    re.match(r"^\d+\.\s+", nxt)):
                break
            paragraph_lines.append(nxt)
            i += 1
        p = doc.add_paragraph()
        add_inline_runs(p, join_wrapped(paragraph_lines))


def audit_document(doc: Document) -> None:
    section = doc.sections[0]
    for margin in (section.left_margin, section.right_margin,
                   section.top_margin, section.bottom_margin):
        assert round(margin.inches, 3) == 1.0
    # All six source tables (execution strata, Stage A, Stage B deltas,
    # erank, probe sensitivity, and correlations) must survive conversion.
    # must survive the conversion.
    assert len(doc.tables) == 6, len(doc.tables)
    for table in doc.tables:
        widths = [int(col.get(f"{{{table._tbl.nsmap['w']}}}w"))
                  for col in table._tbl.tblGrid]
        assert sum(widths) == 9360, sum(widths)


def main() -> None:
    doc = Document()
    configure_styles(doc)
    configure_header_footer(doc)
    doc.core_properties.title = "实验一 结果详解（事实核验修订版·零基础）"
    doc.core_properties.subject = "Algoverse RLVR plasticity project — Experiment 1 fact-checked results"
    doc.core_properties.author = "Aaron Wang research workspace"
    doc.core_properties.keywords = "RLVR, GRPO, plasticity, Experiment 1, 结果解读"
    add_title_block(doc, "实验一 结果详解（事实核验修订版·零基础）")
    add_toc(doc)
    build_body(doc)
    set_update_fields(doc)
    audit_document(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
