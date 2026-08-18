from __future__ import annotations

import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "EXPERIMENT1_COMPLETE_BEGINNER_GUIDE_ZH.md"
OUTPUT = ROOT / "EXPERIMENT1_COMPLETE_BEGINNER_GUIDE_ZH.docx"
ASSET_DIR = ROOT / "document_builders" / "_exp1_assets"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
MUTED = "666666"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
TABLE_BORDER = "C8D2DC"
WHITE = "FFFFFF"
CAUTION = "FFF4CE"
RED_FILL = "FCE8E6"
GREEN_FILL = "E6F4EA"

CJK_FONT = "Arial Unicode MS"
WESTERN_FONT = CJK_FONT
MONO_FONT = "Arial Unicode MS"

FONT_REGULAR_PATH = "/System/Library/Fonts/Hiragino Sans GB.ttc"
FONT_FALLBACK_PATH = "/Library/Fonts/Arial Unicode.ttf"


def set_east_asia_font(rpr, name: str = CJK_FONT) -> None:
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:cs"), name)


def style_font(style, *, western=WESTERN_FONT, east_asia=CJK_FONT, size=None,
               bold=None, color=None) -> None:
    style.font.name = western
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), western)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), western)
    set_east_asia_font(style._element.get_or_add_rPr(), east_asia)
    if size is not None:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if color is not None:
        style.font.color.rgb = RGBColor.from_string(color)


def set_run_font(run, *, western=WESTERN_FONT, east_asia=CJK_FONT, size=None,
                 bold=None, italic=None, color=None) -> None:
    run.font.name = western
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), western)
    rpr.rFonts.set(qn("w:hAnsi"), western)
    set_east_asia_font(rpr, east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_shading(paragraph, fill: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, side: str, color: str, size: int = 8,
                         space: int = 4) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    edge = pbdr.find(qn(f"w:{side}"))
    if edge is None:
        edge = OxmlElement(f"w:{side}")
        pbdr.append(edge)
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), str(space))
    edge.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trpr.append(tbl_header)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tcpr = tc.get_or_add_tcPr()
    tcmar = tcpr.first_child_found_in("w:tcMar")
    if tcmar is None:
        tcmar = OxmlElement("w:tcMar")
        tcpr.append(tcmar)
    for m, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcmar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcmar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    total = sum(widths_dxa)
    tbl = table._tbl
    tblpr = tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(total))
    tblw.set(qn("w:type"), "dxa")

    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(indent_dxa))
    tblind.set(qn("w:type"), "dxa")

    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(width))
            tcw.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)


def set_table_borders(table, color=TABLE_BORDER, size=6) -> None:
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def add_field(paragraph, instr: str, fallback: str = "1") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    inst = OxmlElement("w:instrText")
    inst.set(qn("xml:space"), "preserve")
    inst.text = instr
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = fallback
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, inst, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def add_numbering_definition(doc: Document, *, ordered: bool) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids or [0]) + 1
    num_id = max(num_ids or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    numfmt = OxmlElement("w:numFmt")
    numfmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    lvltext = OxmlElement("w:lvlText")
    lvltext.set(qn("w:val"), "%1." if ordered else "•")
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.extend([tabs, ind, spacing])
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), WESTERN_FONT)
    rfonts.set(qn("w:hAnsi"), WESTERN_FONT)
    rfonts.set(qn("w:eastAsia"), CJK_FONT)
    rpr.append(rfonts)
    lvl.extend([start, numfmt, lvltext, jc, ppr, rpr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    numpr.extend([ilvl, numid])


def add_inline_runs(paragraph, text: str, *, base_size: float | None = None,
                    base_color: str | None = None) -> None:
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=base_size, color=base_color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=base_size, bold=True, color=base_color)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, western=MONO_FONT, east_asia=CJK_FONT,
                         size=(base_size or 11) - 0.5, color=DARK_BLUE)
            set_paragraph_shading(paragraph, "F8FAFC")
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=base_size, color=base_color)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    style_font(normal, size=11, color="111111")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.widow_control = True

    title = doc.styles["Title"]
    # 28 pt keeps the full Chinese/English title on one line in native Word.
    style_font(title, size=28, bold=True, color=NAVY)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.styles["Subtitle"]
    style_font(subtitle, size=14, bold=False, color=DARK_BLUE)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(20)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style_font(style, size=size, bold=True, color=color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for name in ("Caption",):
        style = doc.styles[name]
        style_font(style, size=9, color=MUTED)
        style.font.italic = True
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def configure_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(2)
    run = hp.add_run("Experiment 1 零基础完整入门")
    set_run_font(run, size=9, color=MUTED, bold=True)
    set_paragraph_border(hp, "bottom", "D7DBE2", size=5, space=3)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(3)
    run = fp.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    add_field(fp, "PAGE", "1")
    run = fp.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)


def pil_font(size: int) -> ImageFont.FreeTypeFont:
    path = FONT_REGULAR_PATH if Path(FONT_REGULAR_PATH).exists() else FONT_FALLBACK_PATH
    return ImageFont.truetype(path, size=size)


def pil_box(draw: ImageDraw.ImageDraw, rect: tuple[int, int, int, int], title: str,
            body: str, fill: str, *, title_size=42, body_size=31) -> None:
    draw.rounded_rectangle(rect, radius=18, fill=f"#{fill}", outline="#56758F", width=4)
    x1, y1, x2, y2 = rect
    cx = (x1 + x2) // 2
    cy = (y1 + y2) // 2
    draw.text((cx, cy - 42), title, font=pil_font(title_size), fill="#203748", anchor="mm", align="center")
    draw.multiline_text((cx, cy + 38), body, font=pil_font(body_size), fill="#333333",
                        anchor="mm", align="center", spacing=8)


def pil_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int],
              color="#56758F") -> None:
    draw.line([start, end], fill=color, width=6)
    x2, y2 = end
    x1, y1 = start
    dx, dy = x2 - x1, y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    head = 18
    p1 = (x2 - ux * head + px * 10, y2 - uy * head + py * 10)
    p2 = (x2 - ux * head - px * 10, y2 - uy * head - py * 10)
    draw.polygon([(x2, y2), p1, p2], fill=color)


def make_overview_figure(path: Path) -> None:
    img = Image.new("RGB", (2200, 980), "white")
    draw = ImageDraw.Draw(img)
    draw.text((1100, 82), "研究问题逐步收窄：先校准仪器与剂量，再进入正式 detector bake-off",
              font=pil_font(46), fill="#203748", anchor="mm")
    items = [
        ((55, 350, 405, 620), "正式 Proposal", "预测 Stage-B stall\n并胜过 dashboard", LIGHT_BLUE),
        ((465, 350, 825, 620), "Experiment 1", "小模型 pilot\n验证管线与 RQ1 proxy", "EAF2F8"),
        ((885, 350, 1265, 620), "Experiment 1.5", "剂量升级 + 降噪\n捕获 collapse / 完成长 run", CAUTION),
        ((1325, 180, 1710, 420), "Experiment 1.5.1", "密集 Q 取证\n比较 Q 与 entropy lead", RED_FILL),
        ((1325, 600, 1710, 840), "Experiment 1.6", "3e-6 中间剂量\n待团队批准", GREEN_FILL),
        ((1800, 350, 2140, 620), "Experiment 2", "正式 stall label\n多 run detector", LIGHT_BLUE),
    ]
    for rect, title, body, fill in items:
        pil_box(draw, rect, title, body, fill, title_size=38, body_size=28)
    for a, b in [((410, 485), (455, 485)), ((830, 485), (875, 485)),
                 ((1270, 465), (1315, 310)), ((1270, 505), (1315, 720)),
                 ((1715, 300), (1790, 440)), ((1715, 720), (1790, 535))]:
        pil_arrow(draw, a, b)
    img.save(path)


def make_pipeline_figure(path: Path) -> None:
    img = Image.new("RGB", (2200, 1180), "white")
    draw = ImageDraw.Draw(img)
    draw.text((1100, 78), "Experiment 1 数据流：同一 checkpoint 同时接受内部体检和固定预算体能测试",
              font=pil_font(44), fill="#203748", anchor="mm")
    boxes = [
        ((60, 490, 400, 720), "Qwen2.5-0.5B", "checkpoint 0", LIGHT_BLUE),
        ((500, 490, 900, 720), "Stage A", "GSM8K GRPO\n200 updates", CAUTION),
        ((1040, 175, 1470, 415), "Phase 2：测 Q", "固定 probe\n层 4 / 12 / 22", LIGHT_BLUE),
        ((1040, 490, 1470, 720), "五个 checkpoint", "0 / 25 / 50 / 100 / 200", "EAF2F8"),
        ((1040, 805, 1470, 1045), "Phase 3：适应", "每个起点独立\nSVAMP 50 updates", LIGHT_BLUE),
        ((1720, 490, 2140, 720), "Phase 4：关联", "erank_L12\nvs. svamp_delta", LIGHT_BLUE),
    ]
    for rect, title, body, fill in boxes:
        pil_box(draw, rect, title, body, fill, title_size=40, body_size=30)
    for a, b in [((405, 605), (490, 605)), ((905, 605), (1030, 605)),
                 ((1255, 480), (1255, 425)), ((1255, 730), (1255, 795)),
                 ((1480, 295), (1710, 545)), ((1480, 925), (1710, 665))]:
        pil_arrow(draw, a, b)
    draw.multiline_text((1900, 925), "注意：这不是正式 stall label；\n只是 RQ1 的 checkpoint-level proxy",
                        font=pil_font(28), fill="#9B1C1C", anchor="mm", align="center", spacing=8)
    img.save(path)


def make_dose_figure(path: Path) -> None:
    img = Image.new("RGB", (2100, 700), "white")
    draw = ImageDraw.Draw(img)
    draw.text((1050, 70), "当前剂量地图：目标是找到“不会直接崩，但能留下持久变化”的窗口",
              font=pil_font(45), fill="#203748", anchor="mm")
    y = 335
    draw.line([(250, y), (1850, y)], fill="#9AA0A6", width=8)
    points = [
        (350, "#34A853", "1e-6", "200/500 步健康\n无持久适应退化", "#203748", True),
        (1030, "#1A73E8", "3e-6（1.6）", "中间剂量探针\n待团队批准", "#1A73E8", False),
        (1750, "#EA4335", "1e-5", "≤55 步 policy collapse\n已有真实正样本", "#203748", True),
    ]
    for x, color, title, body, title_color, filled in points:
        if filled:
            draw.ellipse((x-28, y-28, x+28, y+28), fill=color, outline=color, width=5)
        else:
            draw.ellipse((x-28, y-28, x+28, y+28), fill="white", outline=color, width=8)
        draw.text((x, y-100), title, font=pil_font(40), fill=title_color, anchor="mm")
        draw.multiline_text((x, y+135), body, font=pil_font(31), fill="#333333",
                            anchor="mm", align="center", spacing=8)
    img.save(path)


def build_figures() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    figures = {
        "overview": ASSET_DIR / "overview_roadmap.png",
        "pipeline": ASSET_DIR / "pilot_pipeline.png",
        "dose": ASSET_DIR / "dose_map.png",
    }
    make_overview_figure(figures["overview"])
    make_pipeline_figure(figures["pipeline"])
    make_dose_figure(figures["dose"])
    figures["q_actual"] = ROOT / "eaaj-pilot/outputs/exp15_cuda_grpo_gsm8k_c7cc7a1d02d9/analysis/fig_a_q_vs_updates.png"
    figures["delta_actual"] = ROOT / "eaaj-pilot/outputs/exp15_cuda_grpo_gsm8k_c7cc7a1d02d9/analysis/fig_d_delta_by_seed.png"
    return figures


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.25) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    # Give Word and screen readers a useful description instead of an unnamed
    # decorative object.  The visible caption below remains the canonical label.
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", "研究图示")
    cp = doc.add_paragraph(caption, style="Caption")
    cp.paragraph_format.keep_with_next = False


def add_cover(doc: Document) -> None:
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(16)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("RESEARCH FIELD GUIDE · 核查版")
    set_run_font(run, size=10.5, bold=True, color="B07A00")

    p = doc.add_paragraph(style="Title")
    p.add_run("从 Proposal 到 Experiment 1 全部版本")
    p2 = doc.add_paragraph(style="Subtitle")
    p2.add_run("零 ML 基础完整入门")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(54)
    add_inline_runs(sub, "研究问题 · 实验设计 · 所有运行版本 · 真实结果 · 下一步决策", base_size=11, base_color=MUTED)

    lead = doc.add_paragraph()
    lead.paragraph_format.left_indent = Inches(0.42)
    lead.paragraph_format.right_indent = Inches(0.42)
    lead.paragraph_format.space_before = Pt(10)
    lead.paragraph_format.space_after = Pt(42)
    lead.paragraph_format.line_spacing = 1.3
    set_paragraph_shading(lead, LIGHT_GRAY)
    set_paragraph_border(lead, "left", BLUE, size=18, space=6)
    add_inline_runs(
        lead,
        "读完目标：即使没有机器学习背景，也能准确解释 proposal 想检验什么、pilot 实际测了什么、每个版本为何出现，以及哪些结论现在仍然不能说。",
        base_size=11.2,
    )

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(4)
    r = meta.add_run("Aaron Wang · Algoverse AI Research Program")
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta2.add_run("更新至 2026-07-19 · 基于正式 proposal、冻结配置与 run artifacts 核查")
    set_run_font(r, size=9.5, color=MUTED)
    doc.add_page_break()


def add_navigation(doc: Document) -> None:
    h = doc.add_paragraph("阅读导航", style="Heading 1")
    h.paragraph_format.page_break_before = False
    p = doc.add_paragraph()
    add_inline_runs(p, "推荐第一次阅读按 1 → 2 → 5 → 6 → 9 → 12 → 15；需要查术语时直接跳到第十七篇。", base_size=11)
    sections = [
        "第一篇：一页看懂全项目",
        "第二篇：Proposal 从零讲起",
        "第三篇：ML 基础",
        "第四篇：Q 与 dashboard 指标",
        "第五篇：Original Experiment 1 设计",
        "第六篇：早期所有版本",
        "第七篇：Pilot 证据边界",
        "第八篇：Experiment 1.5 的动机",
        "第九篇：Experiment 1.5 v1/v2/v3",
        "第十篇：Experiment 1.5.1",
        "第十一篇：Experiment 1.6",
        "第十二篇：完整版本地图",
        "第十三篇：常见误读",
        "第十四篇：如何检查 run 目录",
        "第十五篇：当前知识边界",
        "第十六篇：零基础 FAQ",
        "第十七篇：术语速查",
        "第十八篇：来源与文档地图",
    ]
    num_id = add_numbering_definition(doc, ordered=True)
    for item in sections:
        p = doc.add_paragraph()
        apply_num(p, num_id)
        add_inline_runs(p, item)
    doc.add_page_break()


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in rows[1]):
        rows.pop(1)
    return rows


def compute_widths(rows: list[list[str]], total: int = 9360) -> list[int]:
    n = max(len(r) for r in rows)
    maxima = []
    for j in range(n):
        lengths = [len(re.sub(r"[*`]", "", r[j])) if j < len(r) else 0 for r in rows]
        maxima.append(max(5, min(max(lengths), 42)))
    min_width = 700 if n >= 6 else 850 if n >= 4 else 1100
    available = total - min_width * n
    if available < 0:
        return [total // n] * (n - 1) + [total - (total // n) * (n - 1)]
    weight_sum = sum(maxima)
    widths = [min_width + int(available * w / weight_sum) for w in maxima]
    widths[-1] += total - sum(widths)
    return widths


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_geometry(table, compute_widths(rows))
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])

    for i, data_row in enumerate(rows):
        for j in range(ncols):
            cell = table.rows[i].cells[j]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if i == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            text = data_row[j] if j < len(data_row) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            if i == 0 or (j > 0 and len(text) <= 12 and not re.search(r"[。；，]", text)):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline_runs(p, text, base_size=9.2)
            if i == 0:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(2)


def add_code_block(doc: Document, code: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.05
    set_paragraph_shading(p, "F3F5F7")
    set_paragraph_border(p, "left", "9AA9B8", size=10, space=5)
    for idx, line in enumerate(code):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, western=MONO_FONT, east_asia=CJK_FONT, size=8.6, color="273746")


def build_body(doc: Document, figures: dict[str, Path]) -> None:
    raw_lines = SOURCE.read_text(encoding="utf-8").splitlines()
    # Skip the source file's title block; the DOCX has a dedicated cover.
    start = next(i for i, line in enumerate(raw_lines) if line.strip() == "---") + 1
    lines = raw_lines[start:]
    i = 0
    bullet_num_id = add_numbering_definition(doc, ordered=False)
    current_ordered_num_id = None
    in_ordered_block = False

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            in_ordered_block = False
            i += 1
            continue
        if stripped == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            set_paragraph_border(p, "bottom", "D7DBE2", size=5, space=2)
            i += 1
            continue
        if stripped.startswith("```"):
            code = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i].rstrip())
                i += 1
            i += 1
            add_code_block(doc, code)
            continue
        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            add_table(doc, parse_table(table_lines))
            in_ordered_block = False
            continue
        heading_match = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            p = doc.add_paragraph(style=f"Heading {level}")
            # Let chapter headings flow naturally.  Forcing every H1 onto a new
            # page created several nearly empty transition pages in native Word.
            if level == 1:
                p.paragraph_format.page_break_before = False
            add_inline_runs(p, text)

            if text == "第一篇：先用一页看懂全项目":
                add_figure(doc, figures["overview"], "图 1　从正式 proposal 到后续实验的关系。1.5.1 与 1.6 是不同问题的两条分支。", 6.3)
            elif text == "5.1 一句话配方":
                add_figure(doc, figures["pipeline"], "图 2　Original Experiment 1 的端到端流程。Q 是内部体检，SVAMP delta 是固定预算体能测试。", 6.3)
            elif text == "9.7 v3 Q 轨迹":
                add_figure(doc, figures["q_actual"], "图 3　Experiment 1.5 v3 的真实 Q 轨迹：L12 早期下探后恢复，L22 持续压缩，dormant fraction 恒为 0。", 6.0)
            elif text == "9.8 v3 的 18 个固定预算适应":
                add_figure(doc, figures["delta_actual"], "图 4　Experiment 1.5 v3 的真实三 seed 适应结果。黑线为三 seed 均值，呈明显 V 型。", 6.0)
            elif text == "11.2 为什么选 3e-6":
                add_figure(doc, figures["dose"], "图 5　当前剂量地图。3e-6 是待团队批准的中间剂量探针，不是已完成结果。", 6.2)
            i += 1
            continue
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.right_indent = Inches(0.15)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            set_paragraph_shading(p, LIGHT_GRAY)
            set_paragraph_border(p, "left", BLUE, size=14, space=5)
            add_inline_runs(p, stripped.lstrip("> "))
            i += 1
            continue
        bullet = re.match(r"^-\s+(.+)$", stripped)
        ordered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if bullet:
            p = doc.add_paragraph()
            apply_num(p, bullet_num_id)
            add_inline_runs(p, bullet.group(1))
            i += 1
            continue
        if ordered:
            if not in_ordered_block:
                current_ordered_num_id = add_numbering_definition(doc, ordered=True)
                in_ordered_block = True
            p = doc.add_paragraph()
            apply_num(p, current_ordered_num_id)
            add_inline_runs(p, ordered.group(1))
            i += 1
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (not nxt or nxt == "---" or nxt.startswith("#") or nxt.startswith("|") or
                    nxt.startswith("```") or nxt.startswith(">") or re.match(r"^-\s+", nxt) or
                    re.match(r"^\d+\.\s+", nxt)):
                break
            paragraph_lines.append(nxt)
            i += 1
        p = doc.add_paragraph()
        add_inline_runs(p, " ".join(paragraph_lines))


def set_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def audit_document(doc: Document) -> None:
    section = doc.sections[0]
    assert round(section.left_margin.inches, 3) == 1.0
    assert round(section.right_margin.inches, 3) == 1.0
    assert round(section.top_margin.inches, 3) == 1.0
    assert round(section.bottom_margin.inches, 3) == 1.0
    assert len(doc.tables) >= 13, len(doc.tables)
    assert len(doc.inline_shapes) >= 5
    for table in doc.tables:
        widths = []
        grid = table._tbl.tblGrid
        for col in grid.findall(qn("w:gridCol")):
            widths.append(int(col.get(qn("w:w"))))
        assert sum(widths) == 9360, sum(widths)
        tblw = table._tbl.tblPr.find(qn("w:tblW"))
        assert tblw is not None and int(tblw.get(qn("w:w"))) == 9360
        tblind = table._tbl.tblPr.find(qn("w:tblInd"))
        assert tblind is not None and int(tblind.get(qn("w:w"))) == 120


def main() -> None:
    figures = build_figures()
    doc = Document()
    configure_styles(doc)
    configure_header_footer(doc)
    doc.core_properties.title = "从 Proposal 到 Experiment 1 全部版本：零 ML 基础完整入门"
    doc.core_properties.subject = "Algoverse RLVR plasticity collapse project beginner guide"
    doc.core_properties.author = "Aaron Wang research workspace"
    doc.core_properties.keywords = "RLVR, GRPO, plasticity, Experiment 1, beginner guide"
    add_cover(doc)
    add_navigation(doc)
    build_body(doc, figures)
    audit_document(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
