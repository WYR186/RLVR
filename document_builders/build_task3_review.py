"""Render the Task-3 four-paper review to .docx.

Same house style as the Experiment-1 documents so the team's Word deliverables
read as one series. The source is bilingual (a Chinese summary section followed
by the English review), so the CJK font override from the Chinese builder is
kept here too.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

import build_exp1_guide as house_style
from build_exp1_guide import (
    BLUE,
    LIGHT_GRAY,
    MUTED,
    NAVY,
    add_code_block,
    add_field,
    add_inline_runs,
    add_numbering_definition,
    add_table,
    apply_num,
    configure_styles,
    parse_table,
    set_paragraph_border,
    set_paragraph_shading,
    set_run_font,
    set_update_fields,
)

# The document mixes Chinese and English; Noto Sans SC + Arial is the pairing
# already validated for the Chinese builds in this directory.
house_style.CJK_FONT = "Noto Sans SC"
house_style.WESTERN_FONT = "Arial"
house_style.MONO_FONT = "Menlo"
house_style.style_font.__kwdefaults__.update(
    western="Arial", east_asia="Noto Sans SC")
house_style.set_run_font.__kwdefaults__.update(
    western="Arial", east_asia="Noto Sans SC")


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "lit review" / "TASK3_FOUR_PAPER_REVIEW.md"
OUTPUT = ROOT / "lit review" / "TASK3_FOUR_PAPER_REVIEW.docx"

TITLE = "Task 3: 四篇可塑性论文源头级评审与 RLVR pipeline 选型"
SUBTITLE = "Algoverse · RLVR Plasticity Project · Person 4 (Aaron) · 2026-07-26"

# Arial Unicode MS lacks glyphs for these; Word renders them as boxes.
GLYPH_REPLACEMENTS = {
    "⚠️ ": "注意：",
    "📌 ": "要点：",
    "⚠️": "注意：",
    "📌": "要点：",
}


def strip_unsupported_glyphs(text: str) -> str:
    for bad, good in GLYPH_REPLACEMENTS.items():
        text = text.replace(bad, good)
    return text


def configure_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(2)
    run = hp.add_run("Task 3 · 四篇必读论文评审 · 源头核实版")
    set_run_font(run, size=9, color=MUTED, bold=True)
    set_paragraph_border(hp, "bottom", "D7DBE2", size=5, space=3)

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(3)
    run = fp.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    add_field(fp, "PAGE", "1")
    run = fp.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph(style="Title")
    add_inline_runs(p, TITLE)
    # The title is long; 20 pt keeps it to two balanced lines.
    for run in p.runs:
        run.font.size = Pt(20)

    p = doc.add_paragraph(style="Subtitle")
    add_inline_runs(p, SUBTITLE)


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("目录 / Contents")
    set_run_font(run, size=13, bold=True, color=NAVY)
    set_paragraph_border(p, "bottom", "D7DBE2", size=5, space=3)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    add_field(p, r'TOC \o "1-2" \h \z \u', "在 Word 中按 F9 更新目录")

    p = doc.add_paragraph()
    run = p.add_run("（在 Word 中打开后，右键目录 → 更新域，即可显示页码。）")
    set_run_font(run, size=9, color=MUTED, italic=True)


def _is_latin_word_char(ch: str) -> bool:
    return ch.isascii() and ch.isalnum()


def join_wrapped(lines: list[str]) -> str:
    text = ""
    for line in lines:
        if not text:
            text = line
        elif _is_latin_word_char(text[-1]) and _is_latin_word_char(line[0]):
            text += " " + line
        else:
            text += line
    return text


def add_callout(doc: Document, lines: list[str]) -> None:
    blocks: list[list[str]] = [[]]
    for line in lines:
        if line:
            blocks[-1].append(line)
        elif blocks[-1]:
            blocks.append([])
    blocks = [b for b in blocks if b]

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    set_paragraph_shading(p, LIGHT_GRAY)
    set_paragraph_border(p, "left", BLUE, size=14, space=5)
    for idx, block in enumerate(blocks):
        if idx:
            p.add_run().add_break()
        add_inline_runs(p, join_wrapped(block))


def build_body(doc: Document) -> int:
    raw = strip_unsupported_glyphs(SOURCE.read_text(encoding="utf-8"))
    lines = raw.splitlines()
    # The single H1 is the title, rendered by add_title_block.
    start = next(i for i, line in enumerate(lines) if line.startswith("# ")) + 1
    lines = lines[start:]

    table_count = 0
    i = 0
    bullet_num_id = add_numbering_definition(doc, ordered=False)
    current_ordered_num_id = None
    in_ordered_block = False

    while i < len(lines):
        raw_line = lines[i]
        stripped = raw_line.strip()
        indent = len(raw_line) - len(raw_line.lstrip())
        if not stripped:
            in_ordered_block = False
            i += 1
            continue

        if stripped == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            set_paragraph_border(p, "bottom", "D7DBE2", size=5, space=2)
            i += 1
            continue

        if stripped.startswith("```"):
            code = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i].rstrip())
                i += 1
            i += 1
            add_code_block(doc, code)
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            add_table(doc, parse_table(table_lines))
            table_count += 1
            in_ordered_block = False
            continue

        heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading:
            level = min(len(heading.group(1)) - 1, 3)
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline_runs(p, heading.group(2))
            i += 1
            continue

        if stripped.startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip().lstrip(">").strip())
                i += 1
            add_callout(doc, quote_lines)
            continue

        bullet = re.match(r"^-\s+(.+)$", stripped)
        ordered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if bullet or ordered:
            item_lines = [(bullet or ordered).group(1)]
            i += 1
            while i < len(lines):
                nxt = lines[i].strip()
                nxt_indent = len(lines[i]) - len(lines[i].lstrip())
                if not nxt or nxt_indent <= indent or re.match(r"^(-|\d+\.)\s+", nxt):
                    break
                item_lines.append(nxt)
                i += 1

            p = doc.add_paragraph()
            if bullet:
                apply_num(p, bullet_num_id)
            else:
                if not in_ordered_block:
                    current_ordered_num_id = add_numbering_definition(doc, ordered=True)
                    in_ordered_block = True
                apply_num(p, current_ordered_num_id)
            if indent:
                p.paragraph_format.left_indent = Inches(0.375 * (1 + indent // 2))
            add_inline_runs(p, join_wrapped(item_lines))
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (not nxt or nxt == "---" or nxt.startswith("#") or nxt.startswith("|") or
                    nxt.startswith("```") or nxt.startswith(">") or re.match(r"^-\s+", nxt) or
                    re.match(r"^\d+\.\s+", nxt)):
                break
            paragraph_lines.append(nxt)
            i += 1
        p = doc.add_paragraph()
        add_inline_runs(p, join_wrapped(paragraph_lines))

    return table_count


def audit_document(doc: Document, source_tables: int) -> None:
    section = doc.sections[0]
    for margin in (section.left_margin, section.right_margin,
                   section.top_margin, section.bottom_margin):
        assert round(margin.inches, 3) == 1.0
    assert len(doc.tables) == source_tables, (len(doc.tables), source_tables)
    for table in doc.tables:
        widths = [int(col.get(f"{{{table._tbl.nsmap['w']}}}w"))
                  for col in table._tbl.tblGrid]
        assert sum(widths) == 9360, sum(widths)
        # A column narrower than ~0.5" cannot hold wrapped 9.2 pt prose.
        assert min(widths) >= 700, (min(widths), len(widths))
    # No markdown emphasis or table pipes may survive into the body text.
    for p in doc.paragraphs:
        text = "".join(r.text for r in p.runs)
        assert "**" not in text, text[:120]


def main() -> None:
    doc = Document()
    configure_styles(doc)
    configure_header_footer(doc)
    doc.core_properties.title = TITLE
    doc.core_properties.subject = (
        "Source-grounded review of four plasticity papers and RLVR pipeline selection")
    doc.core_properties.author = "Aaron Wang research workspace"
    doc.core_properties.keywords = (
        "plasticity, RLVR, GRPO, effective rank, dormant neurons, Plasticine, verl, TRL")
    add_title_block(doc)
    add_toc(doc)
    n_tables = build_body(doc)
    set_update_fields(doc)
    audit_document(doc, n_tables)
    doc.save(OUTPUT)
    print(f"{OUTPUT}  ({n_tables} tables)")


if __name__ == "__main__":
    main()
